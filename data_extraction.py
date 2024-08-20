import os
import pandas as pd
from docx import Document
from docx.table import Table

# 指定word文件所在目录和excel输出文件路径
word_directory = 'yourPath'
output_excel = 'output.xlsx'

# 初始化一个空的 DataFrame 用于存储结果
columns = ['文件名', '机构信息', '测评老师', '姓名', '性别', '出生年月', '测试日期',
           '认知能力_得分', '认知能力_总分', '认知能力_结果',
           '感知觉_得分', '感知觉_总分', '感知觉_结果',
           '注意力_得分', '注意力_总分', '注意力_结果',
           '观察力_得分', '观察力_总分', '观察力_结果',
           '辨识能力_得分', '辨识能力_总分', '辨识能力_结果',
           '排序能力_得分', '排序能力_总分', '排序能力_结果',
           '联想能力_得分', '联想能力_总分', '联想能力_结果']
data = []

def split_score(score):
    if '/' in score:
        return score.split('/')
    return [score, '']

def extract_info_from_table(table: Table):
    info = {}
    info['姓名'] = table.cell(0, 1).text
    info['性别'] = table.cell(0, 5).text
    info['出生年月'] = table.cell(0, 10).text
    info['测试日期'] = table.cell(1, 10).text

    # 提取测评结果并拆分得分
    score_fields = ['认知能力', '感知觉', '注意力', '观察力', '辨识能力', '排序能力', '联想能力']
    cell_positions = [(3, 1), (3, 3), (3, 4), (3, 5), (3, 8), (3, 9), (3, 11)]

    for field, pos in zip(score_fields, cell_positions):
        score1, score2 = split_score(table.cell(*pos).text)
        info[f'{field}_得分'] = score1
        info[f'{field}_总分'] = score2
        info[f'{field}_结果'] = table.cell(pos[0] + 1, pos[1]).text

    return info

def extract_institution_and_teacher(doc):
    """从文档前100个字符中提取机构信息和测评老师"""
    full_text = ""

    for paragraph in doc.paragraphs:
        full_text += paragraph.text.strip()
        if len(full_text) >= 100:  # 提取前100个字符
            break

    first_100_chars = full_text[:100]

    # 分隔提取机构信息和测评老师
    institution_info = ""
    teacher_info = ""

    if "测评老师：" in first_100_chars:
        parts = first_100_chars.split("测评老师：")
        institution_info = parts[0].strip()
        if len(parts[1]) >= 2:  # 确保有足够的字符提取测评老师姓名
            teacher_info = parts[1][:3].strip()  # 只提取三个字符
    else:
        # 如果前100个字符没有找到测评老师，则在全文中搜索
        if "测评老师：" in full_text:
            parts = full_text.split("测评老师：")
            institution_info = parts[0].strip()
            if len(parts[1]) >= 2:  # 确保有足够的字符提取测评老师姓名
                teacher_info = parts[1][:2].strip()  # 只提取两个字符

    return institution_info, teacher_info

# 遍历指定目录中的所有 Word 文件
for filename in os.listdir(word_directory):
    if filename.endswith(".docx"):
        print(f"Processing file: {filename}")
        try:
            # 读取 Word 文件
            doc = Document(os.path.join(word_directory, filename))

            # 提取机构信息和测评老师
            institution_info, teacher_info = extract_institution_and_teacher(doc)

            # 假设信息都在第一个表格中
            if doc.tables:
                table = doc.tables[0]
                info = extract_info_from_table(table)
                info['文件名'] = filename  # 添加文件名
                info['机构信息'] = institution_info  # 添加机构信息
                info['测评老师'] = teacher_info  # 添加测评老师

                # 将提取的信息添加到列表中
                data.append(info)
            else:
                print(f"No tables found in {filename}")
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")

# 将列表转换为 DataFrame
df = pd.DataFrame(data, columns=columns)

# 将 DataFrame 保存为 Excel 文件
df.to_excel(output_excel, index=False)
print(f"信息已成功提取并保存到 {output_excel}")

# 打印一些统计信息
print(f"Processed {len(data)} files.")
print(f"DataFrame shape: {df.shape}")

# 读取保存的Excel文件并检查其内容
df_check = pd.read_excel(output_excel)
print(f"Saved DataFrame shape: {df_check.shape}")
print("First few rows of saved data:")
print(df_check.head())
