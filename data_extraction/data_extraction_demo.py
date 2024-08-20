#用于提取“认知能力”类型测评
#此文件为demo版，与正式版相比没有提取出机构信息和测评老师的信息。
import os
import pandas as pd
import docx
from docx import Document
from docx.table import Table

# 指定word文件所在目录和excel输出文件路径
word_directory = 'yourPath'
output_excel = 'output.xlsx'

# 初始化一个空的 DataFrame 用于存储结果
columns = ['文件名', '测评老师', '姓名', '性别', '出生年月', '测试日期',
           '认知能力_得分', '认知能力_结果',
           '感知觉_得分', '感知觉_结果',
           '注意力_得分', '注意力_结果',
           '观察力_得分', '观察力_结果',
           '辨识能力_得分', '辨识能力_结果',
           '排序能力_得分', '排序能力_结果',
           '联想能力_得分', '联想能力_结果']
data = []

def extract_info_from_table(table: Table):
    info = {}
    info['测评老师'] = table.cell(0, 0).text.split('：')[-1].strip()
    info['姓名'] = table.cell(0, 5).text
    info['性别'] = table.cell(0, 5).text
    info['出生年月'] = table.cell(0, 10).text
    info['测试日期'] = table.cell(1, 10).text

    # 提取测评结果
    info['认知能力_得分'] = table.cell(3, 1).text
    info['认知能力_结果'] = table.cell(4, 1).text
    info['感知觉_得分'] = table.cell(3, 3).text
    info['感知觉_结果'] = table.cell(4, 3).text
    info['注意力_得分'] = table.cell(3, 4).text
    info['注意力_结果'] = table.cell(4, 4).text
    info['观察力_得分'] = table.cell(3, 5).text
    info['观察力_结果'] = table.cell(4, 5).text
    info['辨识能力_得分'] = table.cell(3, 8).text
    info['辨识能力_结果'] = table.cell(4, 8).text
    info['排序能力_得分'] = table.cell(3, 9).text
    info['排序能力_结果'] = table.cell(4, 9).text
    info['联想能力_得分'] = table.cell(3, 11).text
    info['联想能力_结果'] = table.cell(4, 11).text

    return info

# 遍历指定目录中的所有 Word 文件
for filename in os.listdir(word_directory):
    if filename.endswith(".docx"):
        print(f"Processing file: {filename}")
        try:
            # 读取 Word 文件
            doc = Document(os.path.join(word_directory, filename))

            # 假设信息都在第一个表格中
            if doc.tables:
                table = doc.tables[0]
                info = extract_info_from_table(table)
                info['文件名'] = filename  # 添加文件名

                # 将提取的信息添加到列表中
                data.append(info)
            else:
                print(f"No tables found in {filename}")
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")

# 将列表转换为 DataFrame
df = pd.DataFrame(data, columns=columns)

# 将 DataFrame 保存为 Excel 文件
df.to_excel(output_excel, index=False)
print(f"信息已成功提取并保存到 {output_excel}")

# 打印一些统计信息
print(f"Processed {len(data)} files.")
print(f"DataFrame shape: {df.shape}")

# 读取保存的Excel文件并检查其内容
df_check = pd.read_excel(output_excel)
print(f"Saved DataFrame shape: {df_check.shape}")
print("First few rows of saved data:")
print(df_check.head())
