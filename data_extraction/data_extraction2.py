import os
import pandas as pd
from docx import Document
from docx.table import Table

# 指定 word 文件所在目录和 excel 输出文件路径
word_directory = 'D:\桌面\长三院数据\data\ShangHai_YiNuo\言语构音'
output_excel = 'output.xlsx'

# 初始化一个空的 DataFrame 用于存储结果
columns = ['文件名', '机构信息', '测评老师', '姓名', '性别', '出生年月', '测试日期',
           '总得分', '声母得分', '韵母得分', '声调得分']

# 目标音列表
target_sounds = ['包', '泡', '帽', '飞', '刀', '桃', '奶', '辣', '歌', '烤', '喝', '机', '气', '西', '珠', '茶', '书', '热', '字', '草', '扫', '八', '拍', '倒', '糕', '裤', '壳', '蜘', '狮', '枣', '菜', '蛋', '棒', '琴', '冰', '穿', '窗', '蚂', '鹅', '衣', '虾', '脚', '五', '鱼', '姨', '笔', '鸭', '牙', '哑', '讶']

columns.extend(target_sounds)

data = []

def extract_info_from_table(table: Table):
    info = {}
    info['姓名'] = table.cell(0, 1).text
    info['性别'] = table.cell(0, 5).text
    info['出生年月'] = table.cell(0, 10).text
    info['测试日期'] = table.cell(1, 10).text

    # 提取总分和各部分得分
    info['总得分'] = table.cell(3, 1).text.strip('%')
    info['声母得分'] = table.cell(5, 1).text.strip('%')
    info['韵母得分'] = table.cell(5, 3).text.strip('%')
    info['声调得分'] = table.cell(5, 6).text.strip('%')

    # 提取目标音测评结果
    for i, target_sound in enumerate(target_sounds):
        result = table.cell(i + 15, 1).text
        if result == '正确':
            info[target_sound] = 1
        elif result == '错误':
            info[target_sound] = 0
        else:
            info[target_sound] = None
        #info[target_sound] = 1 if result == '正确' else 0

    return info

def extract_institution_and_teacher(doc):
    """从文档前100个字符中提取机构信息和测评老师"""
    full_text = doc.paragraphs[0].text.strip()

    institution_info = ""
    teacher_info = ""

    if "测评老师：" in full_text:
        parts = full_text.split("测评老师：")
        institution_info = parts[0].strip()
        teacher_info = parts[1].strip()

    return institution_info, teacher_info

# 遍历指定目录中的所有 Word 文件
for filename in os.listdir(word_directory):
    if filename.endswith(".docx"):
        print(f"Processing file: {filename}")
        try:
            # 读取 Word 文件
            doc = Document(os.path.join(word_directory, filename))

            # 提取机构信息和测评老师
            institution_info, teacher_info = extract_institution_and_teacher(doc)

            # 假设信息都在第一个表格中
            if doc.tables:
                table = doc.tables[0]
                info = extract_info_from_table(table)
                info['文件名'] = filename  # 添加文件名
                info['机构信息'] = institution_info  # 添加机构信息
                info['测评老师'] = teacher_info  # 添加测评老师

                # 将提取的信息添加到列表中
                data.append(info)
            else:
                print(f"No tables found in {filename}")
        except Exception as e:
            print(f"Error processing {filename}: {str(e)}")

# 将列表转换为 DataFrame
df = pd.DataFrame(data, columns=columns)

# 将 DataFrame 保存为 Excel 文件
df.to_excel(output_excel, index=False)
print(f"信息已成功提取并保存到 {output_excel}")

# 读取保存的 Excel 文件并检查其内容
df_check = pd.read_excel(output_excel)
print(f"Saved DataFrame shape: {df_check.shape}")
print("First few rows of saved data:")
print(df_check.head())